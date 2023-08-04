from docx.oxml import parse_xml
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor, Inches

class DocAcronymMaster:
    """
    A class used to extract text from a Word document, identify acronyms,
    and insert a table of acronyms and their meanings into the document.

    Attributes
    ----------
    doc : Document
        a python-docx Document instance

    Methods
    -------
    get_text()
        Returns the text from the Word document as a string.

    update_document(abbreviations: dict)
        Inserts a table of acronyms and their meanings into the document.
    """

    def __init__(self, doc_path):
        """
        Constructs a new instance of DocAcronymMaster.

        Parameters:
        -----------
        doc_path : str
            The path to the Word document.
        """
        self.doc = Document(doc_path)

    def get_text(self):
        """
        Extracts and returns the text from the Word document as a string.
        """
        return ' '.join([p.text for p in self.doc.paragraphs])


    
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def set_col_widths(table):
        widths = (Inches(0.1), Inches(2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    def update_document(self, abbreviations, path):
        """
        Inserts a table of abbreviations and their definitions into the Word document.

        Parameters:
        -----------
        abbreviations : dict
            A dictionary of abbreviations and their definitions.
        """
        # Ensure that the document has at least two pages
        if len(self.doc.paragraphs):  # Adding paragraphs until we have at least two pages
            self.doc.add_page_break()

        # Create a new paragraph at the start of the second page
        paragraph = self.doc.paragraphs[1]._insert_paragraph_before()  # Insert before the first paragraph
        
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
        
        run = paragraph.add_run('List of Abbreviations')
        font = run.font
        font.color.rgb = RGBColor(173, 216, 230)  # Light blue color

        # Add a table after the new paragraph
        table = self.doc.add_table(rows=1, cols=2)

        # Move table before the selected paragraph
        DocAcronymMaster.move_table_after(table, paragraph)

        # Set table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Abbreviation'
        hdr_cells[1].text = 'Definition'

        # Style the header row
        for cell in table.rows[0].cells:
            paragraph = cell.paragraphs[0]
            # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.font.bold = True
            # Set background color of the header cells
            shading_elm = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="590B0B"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)


        # Add a border to the entire table
        table.style = 'Table Grid'
        # Adjust the row height
        for row in table.rows:
            row.height = Pt(12)  # Adjust as needed

        # Add rows for each abbreviation
        for abbr, defn in abbreviations.items():
            row_cells = table.add_row().cells
            row_cells[0].text = abbr
            row_cells[1].text = defn

            # Adjust row height
            row_cells[0]._element.get_or_add_tcPr().get_or_add_tcW().attrib[qn('w:w')] = '2000'
            row_cells[1]._element.get_or_add_tcPr().get_or_add_tcW().attrib[qn('w:w')] = '4000'


        DocAcronymMaster.set_col_widths(table)


    def saveDocument(self, path):
        # Save the document
        self.doc.save(path)


def main():
    from abbreviation_detector import get_abbreviations

    docMaster = DocAcronymMaster("document.docx")
    # get the text of the document
    text = docMaster.get_text()
    # get the abbreviations in the text
    abbreviations = get_abbreviations(text)
    # update the document with the table of abbreviations
    docMaster.update_document(abbreviations)

if __name__ == "__main__":
    main()
